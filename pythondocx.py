from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob
import os

# Insert an array of chapter names here
chname = ['CH 1 - Hypothesis Testing Using Binomial Distribution', 'CH 2 - Hypothesis Testing Using Normal Distribution', 'CH 3 - Poisson Distribution', 'CH 4 - Linear Combination Of Random Variables', 'CH 5 - Continuous Random Variables', 'CH 6 - Sampling']
ph=''

# Prevent table rows from breaking with page layout
def prevent_document_break(document):
    """https://github.com/python-openxml/python-docx/issues/245#event-621236139
       Globally prevent table cells from splitting across pages.
    """
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag

# Scale the picture width ONLY when it is exceeding the docx page
def scale_picture(picture, new_width):
    aspect_ratio = float(picture.height) / float(picture.width)
    picture.height = int(aspect_ratio * new_width)
    picture.width = new_width

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Loop starts
for chap in range(len(chname)):

    document = Document()

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1.5

    papernum = '7'                        # 2. CHANGE PAPER NUMBER
    paperpath = 'Alevel/MathStats2'       # 3. CHANGE SUBJECT
    QorA = 'A'                            # 4. CHANGE Q or A
    title = chname[chap]
    pic = []
    if QorA == 'Q':
        pic = [f for f in glob.glob("png/"+paperpath+"/Paper"+papernum+"/CH"+str(chap+1)+"/Question/*.png")] 
    else:
        #pic = [f for f in glob.glob("png/"+paperpath+"/Paper"+papernum+"/CH"+str(chap+1)+"/Answer/*.png")]
        pic = [h for h in glob.glob("png/"+paperpath+"/Paper"+papernum+"/CH"+str(chap+1)+"/Answer/*.txt")] ### FOR TXT

    name = []
    for g in range(len(pic)):
        name.append(pic[g].split("\\")[-1][:-4])

    # CREATE TABLE FOR WRITING PNG FILES
    chapter_title = document.add_heading(title, 0)
    title_format = chapter_title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spac = document.add_paragraph().add_run()
    spac.add_break(WD_BREAK.PAGE)

    table = document.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    table.autofit = True
    
    # FOR IMAGES
    """ for i in range(len(pic)):
        if name[i][-2] == '_':
            picture = run.add_picture(pic[i])
            if picture.width > Inches(5.5):
                scale_picture(picture, Inches(5.5))
                ph = ph + (picture.height*Inches(5.5))/picture.width
            else:
                ph = ph + picture.height
            if ph > Inches(8.3):
                print(name[i].split("_")[0], ph)

        else:
            """ document.add_section()
            table = document.add_table(rows=0, cols=1)
            table.style = 'Table Grid'
            table.autofit = True """
            row_cells = table.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.insert_paragraph_before(name[i]+'\n', style='List Number').add_run()
            picture = run.add_picture(pic[i])
            if picture.width > Inches(5.5):
                scale_picture(picture, Inches(5.5))
                ph = (picture.height*Inches(5.5))/picture.width
            else:
                ph = picture.height """ 

    # ONLY FOR OBJECTIVE ANSWERS EXTRACTING
    ans = []
    for j in range(len(pic)):
        with open(pic[j], 'r') as f:
        #file = open(pic[j], encoding="utf8", errors='ignore')
            x = f.readlines()
            ans.append(x)
    
    # CREATE TABLE FOR WRITING TXT FILES
    table = document.add_table(rows=0, cols=5)
    table.style = 'Table Grid'
    table.autofit = True

    colcount = 0
    row_cells = table.add_row().cells
    for i in range(len(pic)):
        paragraph = row_cells[colcount].paragraphs[0]
        run = paragraph.insert_paragraph_before(ans[i], style='List Number').add_run()
        colcount = colcount + 1
        if colcount == 5:
            colcount = 0
            row_cells = table.add_row().cells
   

    # PREVENT BREAK
    prevent_document_break(document)

    # SAVE THE FILE
    document.save('docx/'+paperpath+'/Paper'+papernum+'/'+QorA+'_P'+papernum+'CH'+str(chap+1)+'.docx') ### CHANGE PAPER NUMBER

    if len(pic) == len(name):
        print('YES! docx/'+paperpath+'/Paper'+papernum+'/'+QorA+'_P'+papernum+'CH'+str(chap+1)+'.docx has {} pictures.. -> {}' .format(len(pic),title))
